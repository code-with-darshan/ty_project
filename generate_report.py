"""Generate a formatted Word document for the Indian Legal RAG Assistant project report."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helper functions ──────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(6)
    return p

def mixed(parts):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    return p

def blank():
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.4)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = "Times New Roman"
        r2 = p.add_run(text)
        r2.font.size = Pt(12)
        r2.font.name = "Times New Roman"
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Information Technology")
run.bold = True; run.font.size = Pt(18); run.font.name = "Times New Roman"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Project Report — First Stage\nM.Sc. IT / TY AI/ML (2026–2027)")
r2.font.size = Pt(14); r2.font.name = "Times New Roman"

blank(); blank()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Project Title:")
r3.bold = True; r3.font.size = Pt(13); r3.font.name = "Times New Roman"

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    "Indian Legal RAG Assistant\n"
    "A Domain-Neutral, Evidence-Verified Retrieval-Augmented\n"
    "Generation System for Indian Law"
)
r4.bold = True; r4.font.size = Pt(15); r4.font.name = "Times New Roman"

blank(); blank()

simple_table(
    headers=["", ""],
    rows=[
        ["Submitted By",    "[STUDENT NAME(S)]"],
        ["Roll No.",        "[ROLL NUMBER(S)]"],
        ["Class",           "TY AI/ML"],
        ["Academic Year",   "2026–2027"],
        ["Project Guide",   "[GUIDE NAME], Dept. of IT & AI/ML"],
        ["Department",      "Information Technology & AI/ML"],
        ["Institute",       "[COLLEGE NAME]"],
    ],
    col_widths=[2, 4]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Acknowledgments", 1)

body(
    "We are deeply thankful to Dr. S. K. Singh, Head of the Department of Information "
    "Technology and AI/ML, whose structured project format and departmental guidance gave "
    "us a clear framework within which to develop and present this work."
)
body(
    "Our heartfelt appreciation goes to our project guide, [GUIDE NAME], whose patient "
    "feedback, technical direction, and encouragement kept this project moving forward "
    "through every challenge. Discussions with our guide significantly sharpened our "
    "understanding of RAG pipeline design and AI safety considerations applicable to the "
    "legal domain."
)
body(
    "We extend our thanks to all faculty members of the Department of IT & AI/ML who took "
    "time to offer feedback during informal reviews and whose courses laid the conceptual "
    "groundwork for this project."
)
body(
    "We acknowledge the open-source community whose work made this project technically "
    "viable: the BAAI research group for the bge-small-en-v1.5 embedding model, Microsoft "
    "Research for the MS MARCO MiniLM cross-encoder, Meta AI for the LLaMA 3.1 model "
    "distributed through Ollama, and the ChromaDB team for their lightweight persistent "
    "vector store. Without freely available, high-quality tools such as these, a project of "
    "this scope would require resources far beyond the academic setting."
)
body(
    "A special thanks to our classmates for their thoughtful questions during our informal "
    "presentations, which pushed us to think more carefully about edge cases in the "
    "verification pipeline. And to our families — thank you for accommodating late nights "
    "and project-dominated weekends with patience and humour."
)
blank()
body("[STUDENT NAME(S)]\nTY AI/ML, 2026–2027", bold=True)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_items = [
    ("Acknowledgments", ""),
    ("Table of Contents", ""),
    ("Nomenclature", ""),
    ("Acronyms", ""),
    ("List of Figures", ""),
    ("List of Tables", ""),
    ("1.  Introduction", ""),
    ("    1.1  Background and Motivation", ""),
    ("    1.2  Overview", ""),
    ("    1.3  Research Goals and Approach", ""),
    ("2.  Literature Review", ""),
    ("    2.1  Research Statement (Summary of Literature Review)", ""),
    ("3.  Architecture and Design", ""),
    ("    3.1  Design Strategy", ""),
    ("    3.2  Parametric Analysis", ""),
    ("    3.3  Sensitivity and Uncertainty Analysis", ""),
    ("    3.4  Component-Level Design Details", ""),
    ("4.  Methodology / Algorithm Used and Proposed Solution", ""),
    ("    4.1  Document Ingestion Pipeline", ""),
    ("    4.2  Hybrid Retrieval Algorithm", ""),
    ("    4.3  Query Expansion Module", ""),
    ("    4.4  Reranking Mechanism", ""),
    ("    4.5  Answer Generation Module", ""),
    ("    4.6  Chain-of-Verification Algorithm", ""),
    ("    4.7  Confidence Scoring Algorithm", ""),
    ("    4.8  Safety Guardrail Design", ""),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12); run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(2)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NOMENCLATURE
# ══════════════════════════════════════════════════════════════════════════════
heading("Nomenclature", 1)
simple_table(
    headers=["Symbol / Term", "Definition"],
    rows=[
        ["k", "Number of top results retrieved from the vector store in the initial retrieval stage"],
        ["k'", "Number of results retained after the cross-encoder reranking stage"],
        ["θ", "Relevance threshold; cross-encoder logit score below which a chunk is discarded"],
        ["C(q)", "Set of candidate chunks retrieved for query q"],
        ["R(q)", "Set of retained chunks after reranking"],
        ["V(c)", "Verification ratio; fraction of claims in answer supported by evidence chunks c"],
        ["S", "Composite confidence score; a weighted combination of verification ratio and retrieval quality"],
        ["E(t)", "Dense embedding vector for text t, produced by the bi-encoder model"],
        ["BM25(q, d)", "BM25 sparse relevance score between query q and document d"],
        ["RRF(rank)", "Reciprocal Rank Fusion score for a document at a given rank position"],
    ],
    col_widths=[1.5, 4.5]
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ACRONYMS
# ══════════════════════════════════════════════════════════════════════════════
heading("Acronyms", 1)
simple_table(
    headers=["Acronym", "Full Form"],
    rows=[
        ["RAG","Retrieval-Augmented Generation"],["LLM","Large Language Model"],
        ["NLP","Natural Language Processing"],["BNS","Bharatiya Nyaya Sanhita"],
        ["BNSS","Bharatiya Nagarik Suraksha Sanhita"],["BSA","Bharatiya Sakshya Adhiniyam"],
        ["IPC","Indian Penal Code"],["CrPC","Code of Criminal Procedure"],
        ["IEA","Indian Evidence Act"],["PDF","Portable Document Format"],
        ["API","Application Programming Interface"],["CUDA","Compute Unified Device Architecture"],
        ["CPU","Central Processing Unit"],["GPU","Graphics Processing Unit"],
        ["RRF","Reciprocal Rank Fusion"],["BM25","Best Match 25 (Okapi BM25 ranking function)"],
        ["TF-IDF","Term Frequency–Inverse Document Frequency"],["CoVe","Chain-of-Verification"],
        ["NALSA","National Legal Services Authority"],["RTI","Right to Information"],
        ["IT Act","Information Technology Act"],["IBC","Insolvency and Bankruptcy Code"],
        ["EPF","Employees' Provident Fund"],["GI","Geographical Indications"],
        ["GST","Goods and Services Tax"],["UI","User Interface"],
    ],
    col_widths=[1.2, 4.8]
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES / TABLES
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Figures", 1)
simple_table(
    headers=["Figure No.", "Title"],
    rows=[
        ["Figure 1.1","High-level overview of the Indian Legal RAG Assistant system"],
        ["Figure 3.1","End-to-end system architecture diagram"],
        ["Figure 3.2","Document ingestion pipeline flowchart"],
        ["Figure 3.3","Hybrid retrieval and reranking pipeline"],
        ["Figure 3.4","Verification and confidence scoring flow"],
        ["Figure 4.1","BM25 + Dense Vector hybrid retrieval with RRF fusion"],
        ["Figure 4.2","Chain-of-Verification algorithm flow"],
        ["Figure 4.3","Query expansion using few-shot prompted LLM"],
    ],
    col_widths=[1.5, 4.5]
)

heading("List of Tables", 1)
simple_table(
    headers=["Table No.", "Title"],
    rows=[
        ["Table 3.1","Key system configuration parameters"],
        ["Table 3.2","Supported Indian legal domains and statutes"],
        ["Table 4.1","Confidence score tier thresholds"],
        ["Table 4.2","Safety guardrail refusal conditions"],
    ],
    col_widths=[1.5, 4.5]
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 1: Introduction", 1)
heading("1.1 Background and Motivation", 2)

body(
    "India's legal framework is one of the most voluminous in the world. With upward of "
    "1,500 Central Acts and an even larger body of State legislation — touching every "
    "dimension of life from criminal accountability to corporate governance, from family "
    "disputes to environmental protection — the practical task of finding the law that "
    "applies to one's situation has historically demanded either expensive professional "
    "consultation or years of legal training. For most ordinary Indians, this barrier is "
    "insurmountable. A person who has been cheated, harassed, or harmed often does not "
    "know which Act applies, which section is relevant, or even which authority to approach "
    "first. This is the access-to-justice problem that motivates this project: the gap "
    "between a citizen's legal rights on paper and their ability to understand and exercise "
    "those rights in practice."
)
body(
    "This challenge became substantially more acute in 2023 when the Indian Parliament "
    "enacted three major statutes that collectively replaced colonial-era criminal law: the "
    "Bharatiya Nyaya Sanhita (BNS) 2023 replaced the Indian Penal Code 1860; the Bharatiya "
    "Nagarik Suraksha Sanhita (BNSS) 2023 replaced the Code of Criminal Procedure 1973; "
    "and the Bharatiya Sakshya Adhiniyam (BSA) 2023 replaced the Indian Evidence Act 1872. "
    "These were not minor amendments — they represented a complete renumbering of sections, "
    "introduction of new categories of offences (including organised crime, terrorism, and "
    "cybercrimes), and revised procedural timelines. The practical consequence is that legal "
    "information previously memorised by practitioners, embedded in popular guides, or "
    "indexed by search engines is now potentially outdated or misattributed to wrong section "
    "numbers. A technology solution that works directly from the authoritative statute PDFs "
    "— rather than from potentially stale cached content — is therefore both timely and "
    "necessary."
)
body(
    "Recent advances in Large Language Models (LLMs) and Retrieval-Augmented Generation "
    "(RAG) create a genuine opportunity to address the access-to-justice problem at scale. "
    "The key insight behind RAG is that a language model does not need to memorise all "
    "legal knowledge; instead it can retrieve the relevant portions of a statute at query "
    "time and generate an answer grounded exclusively in that retrieved text. This "
    "distinction is not academic — a pure generative model trained on internet text may "
    "confidently produce a section number, penalty, or procedure that does not exist in "
    "Indian law, because it has interpolated from similar-sounding text in its training "
    "data. In the legal context, such a fabrication is not merely inaccurate; it could "
    "cause a user to miss a legal deadline, approach the wrong forum, or waive a right they "
    "actually have. The entire architectural philosophy of this project is therefore: "
    "retrieve first, verify before answering, and refuse rather than guess."
)
body("Existing legal AI tools in the Indian context suffer from several limitations:")
bullet("Narrow domain coverage — most tools focus exclusively on criminal law or a single statute.")
bullet("No hallucination control — answers are generated without any post-generation verification.")
bullet("No confidence transparency — users receive no indication of how reliable a given answer is.")
bullet("Internet dependency — cloud-based tools raise privacy concerns and are inaccessible in low-bandwidth environments.")
bullet("Poor query understanding — colloquial questions are not translated into formal statutory language before retrieval.")
body(
    "This project directly addresses all five limitations through a novel, locally-hosted, "
    "domain-neutral RAG pipeline enriched with query expansion, hybrid retrieval, "
    "cross-encoder reranking, LLM-generated answers, chain-of-verification (CoVe), and a "
    "three-tier safety guardrail system."
)

heading("1.2 Overview", 2)
body(
    "The Indian Legal RAG Assistant is a locally hosted, privacy-preserving "
    "question-answering system designed to provide evidenced, verified answers about Indian "
    "law to ordinary citizens, law students, and legal professionals. The system operates "
    "entirely on the user's own hardware — no data is transmitted to external servers — "
    "making it suitable for use across a wide variety of environments, including those with "
    "data-sensitivity requirements."
)
body("At its core, the system implements a sophisticated RAG pipeline with the following end-to-end flow:")

steps = [
    ("Knowledge Base Construction: ", "Legal statute PDFs covering more than 40 Indian Acts across 10+ legal domains are loaded, cleaned, chunked, embedded using a bi-encoder model, and stored in a persistent ChromaDB vector database."),
    ("Query Reception and Expansion: ", "When a user poses a query in natural or colloquial language, an LLM-powered Query Expander first translates the query into formal statutory language using few-shot prompting."),
    ("Hybrid Retrieval: ", "The system simultaneously performs dense semantic vector search and sparse BM25 keyword matching. Results are merged using Reciprocal Rank Fusion (RRF)."),
    ("Cross-Encoder Reranking: ", "The top-K fused candidates are re-scored using a cross-encoder model. Candidates below a relevance threshold are filtered out."),
    ("Answer Generation: ", "The filtered evidence chunks are provided as grounded context to a locally running LLaMA 3.1 8B Instruct model (via Ollama)."),
    ("Chain-of-Verification: ", "An independent LLM verification pass decomposes the generated answer into individual claims and checks each claim against the retrieved evidence."),
    ("Confidence Scoring: ", "A composite confidence score is computed as a weighted combination of verification ratio (60%) and normalised retrieval quality (40%)."),
    ("Streamlit Web Interface: ", "The entire pipeline is accessible through an intuitive web-based UI with clickable citations, document upload, and a Personal Incident Support panel."),
]
for i, (bold_part, rest) in enumerate(steps):
    mixed([(f"{i+1}. ", True, False), (bold_part, True, False), (rest, False, False)])

body(
    "The system covers Indian law across criminal (BNS, BNSS, BSA), civil (Indian Contract "
    "Act, Specific Relief Act, Arbitration Act), property, family, consumer, digital, "
    "labour, corporate, intellectual property, environmental, tax, and constitutional law."
)

heading("1.3 Research Goals and Approach", 2)
goals = [
    ("Goal 1 — Hallucination-Free Legal QA: ", "Design a pipeline that structurally prevents the LLM from generating answers not grounded in retrieved statutory text, using Chain-of-Verification."),
    ("Goal 2 — Domain-Neutral Coverage: ", "Build a system that supports any Indian legal Act whose PDF can be ingested, covering 40+ Acts across all major domains."),
    ("Goal 3 — Hybrid Retrieval Superiority: ", "Demonstrate that combining dense vector retrieval with sparse BM25 through Reciprocal Rank Fusion outperforms either method alone."),
    ("Goal 4 — Transparent Confidence: ", "Provide interpretable, calibrated confidence signals (High / Moderate / Refused) so users can gauge reliability."),
    ("Goal 5 — Complete Local Operation: ", "Ensure the entire pipeline runs on the user's own machine without any external API dependency, protecting privacy."),
]
for bold_part, rest in goals:
    mixed([(bold_part, True, False), (rest, False, False)])

body(
    "The research approach combines information retrieval (BM25, dense vector search, RRF, "
    "cross-encoder reranking), NLP (bi-encoder embeddings, LLM query expansion), generative "
    "AI (instruction-tuned LLM for answer generation), and AI safety (Chain-of-Verification, "
    "confidence scoring, multi-layer refusal guards). The system is evaluated against a "
    "curated test set measuring safety guardrail accuracy, confidence score distributions, "
    "and inference latency."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 2: Literature Review", 1)

body(
    "Building a reliable legal question-answering system requires drawing from multiple "
    "research threads simultaneously: how to retrieve relevant text efficiently, how to rank "
    "that text by genuine relevance, how to generate answers grounded strictly in retrieved "
    "evidence, and how to detect and prevent hallucinated claims. The sections below review "
    "the key prior work in each of these areas and explain how that work informed specific "
    "design decisions in this project."
)

lit_sections = [
    (
        "Retrieval-Augmented Generation (RAG):",
        "The RAG paradigm was formalised by Lewis et al. (2020), who showed that a retrieval "
        "component backed by a non-parametric document index could supply a generation model "
        "with factual grounding that its parameters alone could not reliably provide. This was "
        "a significant shift from the 'memorise everything' approach of early language models. "
        "Izacard & Grave (2021) extended this by showing that attending over multiple retrieved "
        "passages simultaneously — their Fusion-in-Decoder architecture — produced measurably "
        "better answers than attending over a single retrieved chunk. Both observations shaped "
        "our design: we retrieve, rerank, and pass the top four highest-quality chunks together "
        "to the generation model."
    ),
    (
        "Bi-Encoder and Cross-Encoder Models:",
        "Sentence-BERT (Reimers & Gurevych, 2019) demonstrated that Siamese bi-encoder "
        "architectures — where query and document are embedded independently and compared via "
        "vector distance — can produce high-quality semantic similarity scores at retrieval speed. "
        "BAAI's bge series, which we adopt, builds on this with contrastive training on "
        "retrieval-specific tasks. Nogueira & Cho (2019) showed that cross-encoder models — "
        "where query and document are concatenated and jointly encoded — achieve dramatically "
        "higher ranking precision. This is precisely the two-stage architecture we employ: fast "
        "bi-encoder for the initial 25-candidate shortlist, precise cross-encoder for the final "
        "reranking to 4 candidates."
    ),
    (
        "Hybrid Retrieval and Reciprocal Rank Fusion:",
        "Neither dense vector search nor sparse keyword search is universally superior — they "
        "are complementary. Dense models excel at semantic matching but may miss passages "
        "containing exact statutory citations such as 'Section 173 BNSS.' BM25 (Robertson et "
        "al., 1994) handles exact keyword matching robustly but struggles with semantic "
        "variation. Multiple IR studies confirm that combining the two consistently beats "
        "either alone. We adopt Reciprocal Rank Fusion (Cormack et al., 2009), which assigns "
        "each document a score based on rank position rather than raw score — making fusion "
        "robust to scale differences between cosine similarity scores and BM25 scores."
    ),
    (
        "Large Language Models for Question Answering:",
        "Meta's LLaMA 3.1 8B Instruct (2024) is a publicly available instruction-tuned model "
        "that demonstrates strong capability at following complex system prompts — a critical "
        "requirement for our use case where we need the model to generate answers strictly "
        "from a provided evidence block. Ollama provides a local inference server with an "
        "OpenAI-compatible API, allowing us to switch underlying models without modifying "
        "application code. Choosing a locally served model eliminates per-query costs, avoids "
        "transmission of sensitive legal queries to third-party servers, and allows operation "
        "in offline environments."
    ),
    (
        "Chain-of-Verification:",
        "Hallucination — where a language model produces confident-sounding but factually "
        "incorrect statements — is well-documented and remains an unsolved problem. In law, it "
        "is a safety hazard. Dhuliawala et al. (2023) proposed Chain-of-Verification (CoVe): "
        "explicitly breaking the answer into individual verifiable claims, then checking each "
        "claim against an authoritative source. We found this approach particularly suited to "
        "legal answers, which naturally decompose into discrete propositions ('the offence is "
        "cognisable', 'the punishment is imprisonment up to seven years', 'bail is "
        "discretionary') — each of which can be cross-referenced against the retrieved section "
        "text. Our AnswerVerifier class operationalises this via a second, independent LLM call."
    ),
    (
        "Legal NLP and Indian Law:",
        "Legal NLP is heavily skewed toward Western legal systems (EUR-Lex, ECHR, US federal "
        "courts). Indian legal NLP has received comparatively little research attention. The "
        "ILDC corpus (Kalamkar et al., 2022) and the FIRE legal track (Malik et al., 2021) "
        "represent steps forward in NER and judgment prediction, but both focus on court "
        "decisions rather than statute text, and neither provides a generative QA interface. "
        "Commercial platforms (SCC Online, Manupatra) offer keyword search but do not "
        "synthesise an answer. The gap this project occupies is: automated, cited, verified, "
        "plain-language answers to statute-level legal questions, produced on-device."
    ),
    (
        "Query Expansion for Legal Retrieval:",
        "A fundamental problem in legal IR is vocabulary mismatch: a user asking 'what happens "
        "if I don't get paid my salary' uses entirely different words from Section 2 of the "
        "Payment of Wages Act. Traditional pseudo-relevance feedback (PRF) is poorly suited "
        "here because if initial retrieval is wrong due to vocabulary mismatch, PRF iterates "
        "on bad evidence. Gao et al. (2022) proposed HyDE (Hypothetical Document Embeddings), "
        "where the LLM generates a hypothetical relevant document for retrieval. We adapt a "
        "related idea: using a few-shot prompted LLM to rewrite the informal query into "
        "formal statutory terminology — producing a rewritten query far more likely to match "
        "statutory text in the vector space."
    ),
]

for title, content in lit_sections:
    mixed([(title, True, False), (" " + content, False, False)])

heading("2.1 Research Statement (Summary of Literature Review)", 2)
body("Taken together, the surveyed literature points to the following concrete design conclusions:")

conclusions = [
    "A retrieval-augmented architecture is non-negotiable for a legal assistant — pure generative approaches carry unacceptable hallucination risk.",
    "Bi-encoder retrieval followed by cross-encoder reranking (retrieve-then-rerank) is the established best practice for retrieval precision.",
    "Combining sparse BM25 with dense vector search and fusing results using RRF yields better coverage than either method in isolation.",
    "Chain-of-Verification applied post-generation is a structurally sound method for catching hallucinated claims before they reach the user.",
    "Indian statute-level QA is an open problem — no publicly documented open-source system provides generative, evidence-verified answers over Indian statute text.",
    "LLM-based query expansion using few-shot examples is more reliable for legal retrieval than traditional feedback methods.",
]
for i, c in enumerate(conclusions):
    bullet(f"{i+1}. {c}")

body(
    "The research gap this project targets is well-defined: no existing publicly available "
    "system combines privacy-preserving local inference, multi-domain Indian statute "
    "coverage, hybrid retrieval with RRF, cross-encoder reranking, Chain-of-Verification, "
    "and a confidence-tiered refusal mechanism in a single deployable application. Each "
    "component has prior art in the literature; their integration in the context of Indian "
    "law constitutes the original contribution of this project."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 3: Architecture and Design", 1)
heading("3.1 Design Strategy", 2)

body(
    "Five explicit design principles governed every architectural decision in this project. "
    "These principles emerged from reflecting on the failure modes of existing legal AI "
    "tools and from the research goals stated in Section 1.3:"
)
principles = [
    ("Principle 1 — Modular Pipeline Architecture: ",
     "The system is composed of six independent, swappable Python classes: Embedder, "
     "VectorSearcher, QueryExpander, the reranker, the generator, and AnswerVerifier. "
     "Each class communicates through well-defined input/output contracts, so that "
     "replacing, say, the embedding model requires changing only one class and one "
     "configuration line."),
    ("Principle 2 — Verification Before Presentation: ",
     "An answer is not delivered to the user unless it has survived the full pipeline: "
     "retrieval, reranking, generation, and independent claim-level verification. In "
     "the legal domain, a hedged-but-wrong answer is more dangerous than a frank refusal."),
    ("Principle 3 — Fully Local Operation: ",
     "Every computation — embedding, BM25 scoring, reranking, LLM inference, and vector "
     "search — runs on the user's local hardware. This was motivated by cost (no per-query "
     "API charges), privacy (legal queries can be sensitive), and resilience (works offline)."),
    ("Principle 4 — Statute-File-Driven Domain Coverage: ",
     "The knowledge base is not statically coded to any particular Acts. Any Indian statute "
     "whose PDF is placed in the data/raw/ directory or uploaded via the web interface is "
     "automatically indexed and labelled. Adding a new legal domain requires no code changes."),
    ("Principle 5 — Actionable Confidence Communication: ",
     "Rather than presenting answers as uniformly authoritative, the system explicitly "
     "reports a confidence score, a confidence tier (High / Moderate), and a per-claim "
     "audit trail, supporting informed decision-making rather than over-reliance."),
]
for bold_part, rest in principles:
    mixed([(bold_part, True, False), (rest, False, False)])

heading("System Architecture Diagram", 2)
arch = (
    "┌───────────────────────────────────────────────────┐\n"
    "│           USER INTERFACE (Streamlit)               │\n"
    "│  [Query Input] [Document Upload] [Citation Viewer] │\n"
    "│  [Confidence Indicator] [Incident Support Panel]   │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "┌───────────────────────────────────────────────────┐\n"
    "│          QUERY PROCESSING LAYER                   │\n"
    "│  QueryExpander (LLaMA 3.1 → legal terminology)   │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "┌───────────────────────────────────────────────────┐\n"
    "│          HYBRID RETRIEVAL LAYER                   │\n"
    "│  Dense Vector Search  +  Sparse BM25 Search       │\n"
    "│              ↘            ↙                       │\n"
    "│              RRF Fusion                           │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "┌───────────────────────────────────────────────────┐\n"
    "│          RERANKING LAYER                          │\n"
    "│  Cross-Encoder (MS MARCO MiniLM-L-6-v2)          │\n"
    "│  Relevance Threshold Filter (θ = -3.0)           │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "┌───────────────────────────────────────────────────┐\n"
    "│          GENERATION LAYER                         │\n"
    "│  LLaMA 3.1 8B Instruct (via Ollama local server) │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "┌───────────────────────────────────────────────────┐\n"
    "│          VERIFICATION LAYER                       │\n"
    "│  AnswerVerifier — Chain-of-Verification           │\n"
    "│  Composite Confidence Score                       │\n"
    "│  Three-tier Safety Guardrail                      │\n"
    "└──────────────────────┬────────────────────────────┘\n"
    "                       ▼\n"
    "         [Verified Answer + Confidence + Citations]"
)
code_block(arch)

heading("3.2 Parametric Analysis", 2)
body("The system exposes several tunable parameters centralised in config/settings.py:")
simple_table(
    headers=["Parameter", "Value", "Role"],
    rows=[
        ["CHUNK_SIZE", "1000 chars", "Maximum size of each text chunk stored in ChromaDB"],
        ["CHUNK_OVERLAP", "200 chars", "Overlap between consecutive chunks to preserve boundary context"],
        ["INITIAL_RETRIEVAL_K", "25", "Candidates retrieved in initial hybrid retrieval stage"],
        ["RERANK_TOP_K", "4", "Chunks retained after cross-encoder reranking"],
        ["RELEVANCE_THRESHOLD (θ)", "−3.0 (logit)", "Score below which a chunk is filtered as insufficiently relevant"],
        ["EMBEDDING_MODEL", "BAAI/bge-small-en-v1.5", "Bi-encoder model for dense embeddings"],
        ["RERANKER_MODEL", "ms-marco-MiniLM-L-6-v2", "Cross-encoder for reranking"],
        ["OLLAMA_MODEL", "llama3.1", "Local LLM for query expansion, generation, and verification"],
        ["RRF_k", "60", "RRF constant controlling rank position influence on fusion score"],
        ["MODEL_DEVICE", "cuda / cpu (auto)", "Hardware selection via PyTorch CUDA availability check"],
        ["CONFIDENCE_REFUSE_THRESHOLD", "40%", "Below this composite score, system refuses to answer"],
        ["Verification ratio weight", "60%", "Weight of claim-level verification ratio in confidence score"],
        ["Retrieval quality weight", "40%", "Weight of normalised reranker score in confidence score"],
    ],
    col_widths=[2, 1.5, 2.5]
)

body("Table 3.2 — Supported Legal Domains and Key Statutes", bold=True)
simple_table(
    headers=["Domain", "Key Acts Supported"],
    rows=[
        ["Criminal Law", "BNS 2023, BNSS 2023, BSA 2023, IPC 1860, CrPC 1973"],
        ["Civil / Contract", "Indian Contract Act 1872, Specific Relief Act 1963, Arbitration Act 1996"],
        ["Property", "Transfer of Property Act 1882, Registration Act 1908"],
        ["Family", "Hindu Marriage Act 1955, Dowry Prohibition Act 1961, PWDV Act 2005"],
        ["Consumer", "Consumer Protection Act 2019"],
        ["Digital / Cyber", "Information Technology Act 2000"],
        ["Labour", "Factories Act 1948, Minimum Wages Act 1948, EPF Act 1952"],
        ["Corporate", "Companies Act 2013, Insolvency and Bankruptcy Code 2016"],
        ["Intellectual Property", "Copyright Act 1957, Patents Act 1970, Trade Marks Act 1999"],
        ["Environmental", "Environment Protection Act 1986, Wildlife Protection Act 1972"],
        ["Tax", "Income Tax Act 1961, GST Act 2017, Customs Act 1962"],
        ["Constitutional", "Constitution of India"],
    ],
    col_widths=[2, 4]
)

heading("3.3 Sensitivity and Uncertainty Analysis", 2)
body(
    "Understanding how the system's behaviour changes as key parameters are varied is "
    "important for maintenance and for understanding safety properties. Three parameters "
    "were found to have the most significant influence on output quality:"
)
sensitivities = [
    ("Sensitivity to INITIAL_RETRIEVAL_K: ",
     "K controls how many candidate chunks are fetched before reranking. K too low (e.g., 5) "
     "risks missing the most relevant chunk. K too high (e.g., 100) increases cross-encoder "
     "latency without improving output quality. K = 25 was found optimal through repeated "
     "testing."),
    ("Sensitivity to the Relevance Threshold (θ): ",
     "θ = −3.0 acts as a hard gate: any chunk with cross-encoder logit below this value is "
     "discarded before generation. Setting θ too high (e.g., 0.0) caused legitimate queries "
     "about niche provisions to be refused. Setting θ too low (e.g., −7.0) allowed tangential "
     "statute passages to reach the generator. θ = −3.0 was selected via evaluation against "
     "a manually labelled set of 15 queries spanning multiple domains."),
    ("Sensitivity to Chunk Size and Overlap: ",
     "Legal statute text has a distinctive structure — provisions are numbered sections, and "
     "the operative legal rule typically appears within the first 800–1200 characters, "
     "followed by explanations and illustrations. CHUNK_SIZE = 1000 characters aligns with "
     "this structure. The 200-character overlap is a safety margin for section boundaries "
     "that fall mid-sentence during PyMuPDF extraction."),
]
for bold_part, rest in sensitivities:
    mixed([(bold_part, True, False), (rest, False, False)])

body("Uncertainty Sources and Mitigations:", bold=True)
simple_table(
    headers=["Uncertainty Source", "Mitigation Strategy"],
    rows=[
        ["LLM hallucination in generation", "Chain-of-Verification: each claim audited against evidence"],
        ["Insufficient evidence in knowledge base", "Composite score threshold: refuse if score < 40%"],
        ["Poor query-document vocabulary match", "LLM-based query expansion to formal statutory language"],
        ["Cross-encoder miscalibration", "Hard relevance threshold filter before generation"],
        ["Scanned / non-machine-readable PDFs", "User-facing warning; system only processes embedded text"],
    ],
    col_widths=[3, 3]
)

heading("3.4 Component-Level Design Details", 2)
components = [
    ("3.4.1 ChromaDB Vector Store: ",
     "ChromaDB is used as the persistent vector store. The collection named 'indian_law' "
     "stores text chunks, their dense embeddings, and metadata (source, act, section, pages, "
     "chunk_index). Persistence is maintained through the PersistentClient, ensuring the "
     "knowledge base survives system restarts without reindexing."),
    ("3.4.2 Embedder (BAAI/bge-small-en-v1.5): ",
     "The bi-encoder uses separate encode_documents() and encode_query() calls, reflecting "
     "the asymmetric query-document embedding approach required by BGE models. Documents are "
     "prefixed with the appropriate instruction for asymmetric retrieval, producing embeddings "
     "optimised specifically for passage retrieval."),
    ("3.4.3 In-Memory BM25 Index: ",
     "On initialisation, VectorSearcher loads the complete document corpus from ChromaDB into "
     "memory and builds a BM25 index using rank-bm25. While memory-intensive, this approach "
     "avoids per-query corpus loading overhead and provides sub-millisecond BM25 scoring."),
    ("3.4.4 Streamlit Web Interface: ",
     "The UI provides: a query text box, a response area displaying the verified answer, a "
     "confidence indicator, an expandable claims audit panel, a citation list with clickable "
     "source PDF links, an expandable document upload section for knowledge base expansion, "
     "and a dynamically-rendered Personal Incident Support card for emergency situations."),
]
for bold_part, rest in components:
    mixed([(bold_part, True, False), (rest, False, False)])
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 4: Methodology / Algorithm Used and Proposed Solution", 1)
heading("4.1 Document Ingestion Pipeline", 2)

body(
    "The build_knowledge_base.py module is responsible for converting raw legal statute PDFs "
    "into a searchable vector store. This is a four-stage sequential pipeline, and the "
    "quality of every downstream retrieval and generation step depends on how well this "
    "stage performs:"
)
ingestion_stages = [
    ("Stage 1 — PDF to Page Text (PyMuPDF): ",
     "The load_pdf() function calls PyMuPDF's fitz.open() to iterate over pages and extract "
     "embedded text, preserving page number metadata. PyMuPDF was selected over alternatives "
     "(pdfminer, pdfplumber, pypdf) after testing on BNS and BNSS PDFs: it produced cleaner "
     "text from multi-column gazette-format layouts."),
    ("Stage 2 — Text Normalisation (clean_text): ",
     "Raw PDF-extracted text carries artefacts that degrade retrieval quality: hyphenated "
     "line-break re-joining, repeated whitespace and blank line compression, removal of "
     "running headers, and Unicode normalisation for transliterated Hindi legal terms."),
    ("Stage 3 — Section-Aware Chunking (chunk_pages): ",
     "Rather than slicing text at a fixed character count, chunk_pages() first identifies "
     "section boundaries using regular expressions matching Indian statute formatting patterns "
     "(e.g., ^\\d+\\.\\s, ^Section\\s+\\d+). This section-alignment means most chunks begin "
     "at a section boundary and contain a semantically complete legal provision."),
    ("Stage 4 — Embedding and Persistent Storage: ",
     "Chunks are batch-processed through the BGE bi-encoder's encode_documents() method. The "
     "resulting float32 embedding vectors are stored in ChromaDB via collection.upsert(), "
     "which performs an insert-or-update based on chunk ID, preventing duplicates."),
]
for bold_part, rest in ingestion_stages:
    mixed([(bold_part, True, False), (rest, False, False)])

code_block(
    "For each PDF in data/raw/:\n"
    "   1. load_pdf()  → list of {page_number, text}\n"
    "   2. clean_text() → normalised text per page\n"
    "   3. chunk_pages() → list of {text, act, section, pages, chunk_index}\n"
    "   4. embedder.encode_documents() → dense vectors\n"
    "   5. collection.upsert(ids, documents, embeddings, metadatas)"
)

heading("4.2 Hybrid Retrieval Algorithm", 2)
body("The VectorSearcher.search() method implements a three-stage hybrid retrieval algorithm:")
retrieval_stages = [
    ("Stage 1 — Dense Vector Search: ",
     "The expanded query is embedded using embedder.encode_query(). The resulting query "
     "vector is compared against all stored chunk embeddings in ChromaDB using cosine "
     "similarity. The top-K (= 25) closest chunks are retrieved."),
    ("Stage 2 — Sparse BM25 Search: ",
     "The query is tokenised and scored against the in-memory BM25 index using the "
     "Okapi BM25 formula: BM25(q,d) = Σᵢ IDF(qᵢ) · [f(qᵢ,d) · (k₁+1)] / "
     "[f(qᵢ,d) + k₁ · (1−b+b·|d|/avgdl)]. The top-K BM25-scored chunks are identified."),
    ("Stage 3 — Reciprocal Rank Fusion (RRF): ",
     "Chunks from both stages are merged using RRF. For each chunk at rank r: "
     "RRF(r) = 1 / (k + r), where k = 60. A chunk appearing in both result lists receives "
     "contributions from both. The merged list is sorted by descending total RRF score."),
]
for bold_part, rest in retrieval_stages:
    mixed([(bold_part, True, False), (rest, False, False)])

heading("4.3 Query Expansion Module", 2)
body(
    "The QueryExpander.expand_query() method uses the locally running LLaMA 3.1 model with "
    "a carefully designed few-shot system prompt to translate colloquial queries into formal "
    "Indian legal terminology. The system prompt provides 10 diverse examples covering "
    "criminal, civil, family, consumer, digital, and other legal domains."
)
body("Key Design Decisions:", bold=True)
bullet("temperature=0.0 ensures deterministic, reproducible expansion.")
bullet("The system prompt explicitly prohibits the model from answering the question — it must only reformulate the query.")
bullet("A graceful fallback returns the original query unmodified if expansion fails.")

body("Example query expansions:", bold=True)
simple_table(
    headers=["Original Query", "Expanded Query"],
    rows=[
        ["What happens if someone steals my phone?", "theft of movable property dishonestly taking without consent punishment Bharatiya Nyaya Sanhita BNS"],
        ["Can I break a contract if they lied to me?", "voidability of contract induced by misrepresentation fraud Indian Contract Act 1872"],
    ],
    col_widths=[2.5, 3.5]
)

heading("4.4 Reranking Mechanism", 2)
body(
    "Where the bi-encoder produces independent representations of query and document and "
    "compares them via cosine distance, the cross-encoder (cross-encoder/ms-marco-MiniLM-L-6-v2) "
    "takes the concatenation [CLS] query [SEP] document [SEP] as input and produces a single "
    "scalar relevance score. Because the self-attention mechanism sees both texts together, "
    "the model can identify fine-grained relevance signals in a way that independent "
    "embeddings cannot. Cross-encoding 25 pairs takes approximately 300–800 milliseconds on "
    "CPU — restricting the cross-encoder to only the top-25 RRF candidates keeps latency "
    "acceptable."
)
body("Reranking Procedure:", bold=True)
for step in [
    "Create (query, chunk_text) pairs for each of the top-25 RRF candidates.",
    "Batch-score all pairs using the cross-encoder.",
    "Sort candidates by descending cross-encoder logit score.",
    "Filter out any chunk with logit score < θ = −3.0 (hard relevance threshold).",
    "Return the top-4 (RERANK_TOP_K) surviving chunks.",
]:
    bullet(step)
body(
    "If fewer than 1 chunk survives the threshold, the system refuses to generate an answer, "
    "triggering the 'Insufficient Evidence' safety response."
)

heading("4.5 Answer Generation Module", 2)
body(
    "The answer generation module sends a structured prompt to the LLaMA 3.1 8B Instruct "
    "model running locally via Ollama. temperature=0.0 is used to maximise factual "
    "consistency and reproducibility. The instruction to cite section numbers enables the "
    "UI to link answers back to source PDFs."
)
code_block(
    "SYSTEM: You are an expert Indian legal assistant. Answer ONLY from the\n"
    "provided evidence chunks. Cite the relevant section numbers. If the\n"
    "evidence does not support the question, say so explicitly.\n\n"
    "USER: QUERY: {expanded_query}\n\n"
    "EVIDENCE:\n"
    "[Chunk 1 — {act}, Section {section}]: {chunk_text}\n"
    "[Chunk 2 — {act}, Section {section}]: {chunk_text}\n"
    "...\n"
    "Please provide a clear, accurate answer citing relevant sections."
)

heading("4.6 Chain-of-Verification Algorithm", 2)
body("The AnswerVerifier.verify_and_score() method implements a structured CoVe protocol:")
code_block(
    "Algorithm: Chain-of-Verification\n\n"
    "Input:  query q, generated_answer a, evidence_chunks C\n"
    "Output: verified_answer, confidence_score, confidence_tier, status\n\n"
    "1.  IF a begins with 'I cannot answer' OR C is empty:\n"
    "        RETURN status='refused', confidence=0.0\n\n"
    "2.  Construct verification prompt:\n"
    "    - System: 'Decompose the DRAFT ANSWER into discrete claims.\n"
    "               Check each claim against the PROVIDED EVIDENCE.\n"
    "               Output JSON: {claims:[{claim:..., supported:true/false}]}\n"
    "               DO NOT use outside knowledge.'\n\n"
    "3.  Call LLaMA 3.1 with temperature=0.0, response_format=json_object\n\n"
    "4.  Parse JSON → claims_list\n\n"
    "5.  verification_ratio = supported_count / total_count\n\n"
    "6.  normalised_rerank = clamp((avg_rerank_score + 1.0) / 2.0, 0, 1)\n\n"
    "7.  composite_score = (verification_ratio×0.60 + normalised_rerank×0.40) × 100\n\n"
    "8.  IF composite_score < 40:    RETURN 'Insufficient Evidence'\n"
    "9.  IF any claim unsupported:   RETURN 'Hallucination Detected'\n"
    "10. IF score >= 80: tier='High Confidence'\n"
    "    ELSE:           tier='Moderate Confidence'\n"
    "11. RETURN verified_answer, confidence_score, tier, status='verified'"
)

heading("4.7 Confidence Scoring Algorithm", 2)
body("The composite confidence score S is computed as:")
body("S = (V × 0.60 + N_rerank × 0.40) × 100", bold=True)
body("where:")
bullet("V = verification ratio = (number of supported claims) / (total claims)")
bullet("N_rerank = normalised reranker score = clamp((avg_rerank_score + 1.0) / 2.0, 0.0, 1.0)")
body(
    "This weighting (60% verification quality, 40% retrieval quality) reflects the primacy "
    "of claim-level verification in the legal domain, where factual accuracy is more "
    "important than retrieval recall."
)
body("Table 4.1 — Confidence Score Tier Thresholds", bold=True)
simple_table(
    headers=["Score Range", "Tier", "Meaning"],
    rows=[
        ["S ≥ 80%", "High Confidence", "All claims supported; strong retrieval quality"],
        ["40% ≤ S < 80%", "Moderate Confidence", "All claims supported; moderate retrieval quality"],
        ["S < 40%", "Refused", "Evidence insufficient; no answer presented"],
        ["Any unsupported claim", "Refused (Hallucination)", "One or more claims not grounded in evidence"],
    ],
    col_widths=[1.5, 2, 2.5]
)

heading("4.8 Safety Guardrail Design", 2)
body("The system implements a four-layer safety architecture operating at different pipeline stages:")
body("Table 4.2 — Safety Guardrail Refusal Conditions", bold=True)
simple_table(
    headers=["Layer", "Trigger", "Response"],
    rows=[
        ["Layer 1 — Retrieval Guard", "No chunks survive the relevance threshold", "Insufficient Evidence — No Answer Generated"],
        ["Layer 2 — Verification Score Guard", "Composite confidence score < 40%", "Insufficient Evidence — No Answer Generated"],
        ["Layer 3 — Hallucination Guard", "One or more claims not supported by evidence", "Hallucination Detected — No Answer Generated"],
        ["Layer 4 — Service Unavailability", "LLM or verifier API call fails", "Verification Unavailable — No Answer Generated"],
    ],
    col_widths=[2, 2, 2]
)
body(
    "A deliberate design decision was made to make each refusal informative rather than "
    "generic. Instead of returning a single error message, the system returns one of four "
    "distinct refusal messages, each explaining precisely why the system chose not to answer "
    "and what the user can do next. This specificity was found during testing to substantially "
    "reduce user frustration in scenarios where a legitimate query triggered a refusal due "
    "to an overly broad or informally phrased question."
)
body(
    "The Personal Incident Support feature was implemented as a response to a use-case that "
    "emerged during early system testing: users who described being a victim of a crime were "
    "receiving purely legal information when they actually needed immediate safety guidance "
    "first. When triggered, a deterministic guidance card is rendered using hardcoded official "
    "information: the national emergency number (112), the cyber fraud helpline (1930), the "
    "National Cyber Crime Reporting Portal URL, evidence preservation steps, and NALSA free "
    "legal aid contacts (15100). No part of this guidance panel is generated by the LLM — "
    "every phone number, URL, and procedural step is hardcoded and verified against official "
    "government sources."
)

page_break()
body(
    "Note to Examiner: The subsequent chapters (Chapter 5: Validation of Modelling Technique, "
    "Chapter 6: Performance Evaluation Results and Discussion, Chapter 7: Conclusions and "
    "Recommendations, Chapter 8: References, Chapter 9: Published Paper) will be submitted "
    "as part of the second-stage project documentation upon completion of the full system "
    "evaluation cycle.",
    italic=True
)
blank()
body("[STUDENT NAME(S)]\nRoll No.: [ROLL NUMBER(S)]\nTY AI/ML, 2026–2027\nDepartment of Information Technology & AI/ML\n[COLLEGE NAME]", bold=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"d:\Legal-rag-assistant\Project_Report_Stage1.docx"
doc.save(out_path)
print(f"✅ Document saved: {out_path}")
