"""
Convert System_Analysis_and_Design.md to a properly formatted DOCX.
Run: python generate_sad_docx.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_heading_style(para, level: int):
    """Apply built-in Heading styles."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    para.style = style_map.get(level, "Heading 4")


def add_horizontal_rule(doc):
    """Add a thin horizontal line (mimics '---')."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_code_block(doc, code_text: str):
    """Add a code / diagram block with a shaded background."""
    para = doc.add_paragraph()
    para.style = "Normal"
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

    # Shade the paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)

    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)


def add_markdown_paragraph(doc, line: str):
    """Handle inline bold/code in a paragraph line."""
    para = doc.add_paragraph()
    para.style = "Normal"
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    _apply_inline(para, line)
    return para


def add_bullet(doc, line: str, indent: int = 0):
    """Add a bullet list item."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    _apply_inline(para, line)


def _apply_inline(para, text: str):
    """Parse **bold** and `code` inline markers."""
    # Combined regex for **bold** and `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(0).startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
        else:
            run = para.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def add_table(doc, headers: list, rows: list):
    """Add a formatted table."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header fill
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            p = cell.paragraphs[0]
            p.clear()
            _apply_inline(p, str(val))
            # Alternate row shading
            if r_i % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8EFF7')
                tcPr.append(shd)
            for run in p.runs:
                run.font.size = Pt(9)

    # Column widths — distribute evenly
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(6.5 / col_count)

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Title Page
# ---------------------------------------------------------------------------

def create_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("System Analysis and Design")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Nyaya — Indian Legal RAG Assistant")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    meta_lines = [
        ("Project:", "Indian Legal Knowledge System"),
        ("Document Type:", "System Analysis and Design Report"),
        ("Prepared for:", "Software Projects — College Submission"),
        ("Date:", "August 2026"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_and_build(md_path: str, out_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    create_title_page(doc)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n").rstrip("\r")

        # --- Code block (mermaid / python / etc.) ---
        if line.startswith("```"):
            if not in_code:
                in_code = True
                lang = line[3:].strip()
                code_lines = []
                if lang:
                    code_lines.append(f"[{lang.upper()} DIAGRAM / CODE BLOCK]")
            else:
                in_code = False
                block_text = "\n".join(code_lines)
                add_code_block(doc, block_text)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # --- Headings ---
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            para = doc.add_paragraph(text)
            set_heading_style(para, level)
            i += 1
            continue

        # --- Table detection ---
        if line.startswith("|"):
            # Collect table lines
            tbl_lines = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                tbl_lines.append(lines[i].rstrip())
                i += 1
            # Parse
            if len(tbl_lines) >= 2:
                hdrs = [c.strip() for c in tbl_lines[0].strip("|").split("|")]
                data_start = 2 if (len(tbl_lines) > 1 and re.match(r'^\|[\s\-|:]+\|$', tbl_lines[1])) else 1
                rows_data = []
                n_cols = len(hdrs)
                for tl in tbl_lines[data_start:]:
                    row = [c.strip() for c in tl.strip("|").split("|")]
                    # Pad or trim to match header column count
                    row = (row + [""] * n_cols)[:n_cols]
                    rows_data.append(row)
                add_table(doc, hdrs, rows_data)
            continue

        # --- Bullet points ---
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            add_bullet(doc, bullet_match.group(3), indent)
            i += 1
            continue

        # --- Numbered list ---
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if numbered_match:
            indent = len(numbered_match.group(1)) // 2
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            _apply_inline(para, numbered_match.group(2))
            i += 1
            continue

        # --- Empty line (spacing) ---
        if line.strip() == "":
            i += 1
            continue

        # --- Normal paragraph ---
        add_markdown_paragraph(doc, line)
        i += 1

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    MD_FILE  = r"d:\Legal-rag-assistant\System_Analysis_and_Design.md"
    OUT_FILE = r"d:\Legal-rag-assistant\System_Analysis_and_Design.docx"
    parse_and_build(MD_FILE, OUT_FILE)
